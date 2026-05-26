"""
comparison_builder.py — Xây dựng bảng so sánh nội dung cũ/mới.
Xuất ra DOCX (A4 ngang) và XLSX.
"""

import re
from datetime import datetime

from .models import ComparisonRow


_DOC_TYPE_PREFIX = re.compile(
    r'^(?:Thông tư liên tịch|Thông tư|Nghị định|Quyết định|'
    r'Luật|Pháp lệnh|Chỉ thị|Bộ luật|Hiến pháp|Văn bản)\s+',
    re.IGNORECASE
)

def _short_title(title: str) -> str:
    return _DOC_TYPE_PREFIX.sub('', title).strip()


# ── Màu sắc theo loại thao tác ────────────────────────────
_OP_COLORS = {
    "Sửa đổi"          : {"row_fill": "FFF2CC", "label_fill": "F4B942", "label_text": "000000"},
    "Bổ sung"          : {"row_fill": "E2EFDA", "label_fill": "70AD47", "label_text": "FFFFFF"},
    "Bãi bỏ"           : {"row_fill": "FCE4D6", "label_fill": "C00000", "label_text": "FFFFFF"},
    "Thay cụm từ"      : {"row_fill": "DDEBF7", "label_fill": "2E75B6", "label_text": "FFFFFF"},
    "Sửa đổi, bổ sung" : {"row_fill": "EAE0F0", "label_fill": "7030A0", "label_text": "FFFFFF"},
    "Đổi tên"          : {"row_fill": "EDEDED", "label_fill": "595959", "label_text": "FFFFFF"},
}
_COLOR_OLD_PHRASE = "C00000"   # đỏ đậm — cụm từ bị xoá
_COLOR_NEW_PHRASE = "375623"   # xanh lá đậm — cụm từ thay thế
_COLOR_HEADER_BG  = "1F4E79"   # xanh navy — header bảng


class ComparisonTableBuilder:
    """
    Xây dựng bảng so sánh nội dung cũ/mới từ merge_log.
    Xuất ra DOCX (file riêng) và XLSX (file riêng).
    """

    OP_LABEL_MAP = {
        "SỬA ĐỔI"          : "Sửa đổi",
        "BỔ SUNG"          : "Bổ sung",
        "BÃI BỎ"           : "Bãi bỏ",
        "THAY CỤM TỪ"      : "Thay cụm từ",
        "SỬA ĐỔI, BỔ SUNG" : "Sửa đổi, bổ sung",
        "ĐỔI TÊN"          : "Đổi tên",
    }
    HEADERS         = ["STT", "Loại thao tác", "Phạm vi tác động",
                        "Nội dung cũ", "Nội dung mới", "Tham chiếu"]
    COL_WIDTHS_CM   = [1.2,  2.8,  5.5,  7.5,  7.5,  5.5]
    COL_WIDTHS_XLSX = [6,    16,   30,   42,   42,   30]

    def __init__(self, merge_log: list, meta: dict):
        self.merge_log = merge_log
        self.meta      = meta
        self.rows      = self._build_rows()

    # ── xây dữ liệu ──────────────────────────────────────

    def _build_rows(self) -> list:
        rows, stt = [], 0
        for entry in self.merge_log:
            action = entry.get("action", "")
            if action == "CẢNH BÁO":
                continue
            stt += 1
            rows.append(ComparisonRow(
                stt           = stt,
                operation     = self.OP_LABEL_MAP.get(action, action),
                scope         = entry.get("node_title") or entry.get("target_scope", ""),
                old_content   = entry.get("original_content", ""),
                new_content   = entry.get("new_content", ""),
                citation      = entry.get("citation", "").strip("()").strip(),
                highlight_old = entry.get("from", ""),
                highlight_new = entry.get("to",   ""),
            ))
        return rows

    # ── DOCX ──────────────────────────────────────────────

    def write_docx(self, output_path: str):
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        for section in doc.sections:
            section.page_width    = Cm(29.7)
            section.page_height   = Cm(21.0)
            section.left_margin   = section.right_margin  = Cm(1.5)
            section.top_margin    = section.bottom_margin = Cm(1.5)
            section.orientation   = 1

        t = doc.add_heading("BẢNG SO SÁNH NỘI DUNG THAY ĐỔI", level=1)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in t.runs:
            run.font.size      = Pt(14)
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        for line in [
            f"Văn bản gốc      : {_short_title(self.meta.get('base_doc', ''))}",
            f"Văn bản sửa đổi  : {', '.join(_short_title(d) for d in self.meta.get('amendment_docs', []))}",
            f"Ngày lập bảng : {self.meta.get('date', datetime.now().strftime('%d/%m/%Y'))}",
            f"Tổng số thay đổi: {len(self.rows)}",
        ]:
            p = doc.add_paragraph(line)
            if p.runs: p.runs[0].font.size = Pt(10)
        doc.add_paragraph()

        dxa   = [int(w * 567) for w in self.COL_WIDTHS_CM]
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        self._docx_set_table_width(table, sum(dxa))

        hdr = table.rows[0]
        self._docx_set_row_widths(hdr, dxa)
        for cell, text in zip(hdr.cells, self.HEADERS):
            self._docx_write(cell, text, Pt(9), bold=True,
                             text_color="FFFFFF", bg=_COLOR_HEADER_BG,
                             align=WD_ALIGN_PARAGRAPH.CENTER)

        for row in self.rows:
            c  = _OP_COLORS.get(row.operation, _OP_COLORS["Sửa đổi"])
            tr = table.add_row()
            self._docx_set_row_widths(tr, dxa)
            cells = tr.cells

            self._docx_write(cells[0], str(row.stt), Pt(9),
                             align=WD_ALIGN_PARAGRAPH.CENTER, bg=c["row_fill"])
            self._docx_write_badge(cells[1], row.operation,
                                   c["label_fill"], c["label_text"])
            self._docx_write(cells[2], row.scope, Pt(9), bg=c["row_fill"])
            self._docx_write_highlight(
                cells[3], row.old_content,
                row.highlight_old, _COLOR_OLD_PHRASE,
                bg=c["row_fill"],
                strike=(row.operation == "Bãi bỏ"))
            self._docx_write_highlight(
                cells[4], row.new_content,
                row.highlight_new, _COLOR_NEW_PHRASE,
                bg=c["row_fill"])
            self._docx_write(cells[5], row.citation, Pt(8.5),
                             italic=True, bg=c["row_fill"])

        doc.save(output_path)

    # ── DOCX helpers ──────────────────────────────────────

    def _get_or_create_tcPr(self, cell):
        """Lấy hoặc tạo mới phần tử w:tcPr của một ô bảng."""
        from docx.oxml.ns import qn
        from docx.oxml   import OxmlElement
        tc   = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        return tcPr

    def _docx_set_table_width(self, table, total_dxa: int):
        from docx.oxml.ns import qn
        from docx.oxml   import OxmlElement
        tbl   = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        w = OxmlElement("w:tblW")
        w.set(qn("w:w"), str(total_dxa))
        w.set(qn("w:type"), "dxa")
        tblPr.append(w)

    def _docx_set_row_widths(self, row, dxa_list: list):
        from docx.oxml.ns import qn
        from docx.oxml   import OxmlElement
        for cell, dxa in zip(row.cells, dxa_list):
            tcPr = self._get_or_create_tcPr(cell)
            w = OxmlElement("w:tcW")
            w.set(qn("w:w"), str(dxa))
            w.set(qn("w:type"), "dxa")
            tcPr.append(w)

    def _docx_set_bg(self, cell, hex_color: str):
        from docx.oxml.ns import qn
        from docx.oxml   import OxmlElement
        tcPr = self._get_or_create_tcPr(cell)
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def _docx_set_margins(self, cell, top=60, bottom=60, left=100, right=100):
        from docx.oxml.ns import qn
        from docx.oxml   import OxmlElement
        tcPr = self._get_or_create_tcPr(cell)
        mar  = OxmlElement("w:tcMar")
        for side, val in [("top", top), ("bottom", bottom),
                          ("left", left), ("right", right)]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"), str(val))
            el.set(qn("w:type"), "dxa")
            mar.append(el)
        tcPr.append(mar)

    def _rgb(self, hex6: str):
        from docx.shared import RGBColor
        return RGBColor(int(hex6[:2], 16), int(hex6[2:4], 16), int(hex6[4:], 16))

    def _docx_write(self, cell, text: str, size,
                    bold=False, italic=False, strike=False,
                    text_color="000000", bg="",
                    align=None):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p   = cell.paragraphs[0]
        if align: p.alignment = align
        run = p.add_run(text or "")
        run.font.size      = size
        run.bold           = bold
        run.font.italic    = italic
        run.font.strike    = strike
        run.font.color.rgb = self._rgb(text_color)
        self._docx_set_margins(cell)
        if bg: self._docx_set_bg(cell, bg)

    def _docx_write_badge(self, cell, text: str, fill_hex: str, txt_hex: str):
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold           = True
        run.font.size      = Pt(8.5)
        run.font.color.rgb = self._rgb(txt_hex)
        self._docx_set_bg(cell, fill_hex)
        self._docx_set_margins(cell)

    def _docx_write_highlight(self, cell, text: str,
                               phrase: str, phrase_color: str,
                               bg: str = "", strike: bool = False):
        from docx.shared import Pt
        p = cell.paragraphs[0]
        if phrase and phrase in (text or ""):
            parts = (text or "").split(phrase)
            for i, part in enumerate(parts):
                if part:
                    r = p.add_run(part)
                    r.font.size   = Pt(9)
                    r.font.strike = strike
                if i < len(parts) - 1:
                    r2 = p.add_run(phrase)
                    r2.font.size      = Pt(9)
                    r2.bold           = True
                    r2.font.color.rgb = self._rgb(phrase_color)
                    r2.font.strike    = strike
        else:
            r = p.add_run(text or "")
            r.font.size   = Pt(9)
            r.font.strike = strike
        self._docx_set_margins(cell)
        if bg: self._docx_set_bg(cell, bg)

    # ── XLSX ──────────────────────────────────────────────

    def write_xlsx(self, output_path: str):
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils  import get_column_letter

        wb = Workbook()
        ws = wb.active
        ws.title = "Bảng so sánh"

        thin  = Side(style="thin",   color="AAAAAA")
        thick = Side(style="medium", color="1F4E79")

        meta_rows = [
            "BẢNG SO SÁNH NỘI DUNG THAY ĐỔI",
            (f"Văn bản gốc: {_short_title(self.meta.get('base_doc', ''))}  |  "
             f"Văn bản sửa đổi: {', '.join(_short_title(d) for d in self.meta.get('amendment_docs', []))}"),
            (f"Ngày lập: {self.meta.get('date', datetime.now().strftime('%d/%m/%Y'))}"
             f"  |  Tổng số thay đổi: {len(self.rows)}"),
        ]
        for i, txt in enumerate(meta_rows, 1):
            ws.merge_cells(f"A{i}:F{i}")
            c = ws[f"A{i}"]
            c.value     = txt
            c.font      = Font(name="Arial", bold=(i == 1),
                               size=13 if i == 1 else 10,
                               color="1F4E79" if i == 1 else "444444")
            c.alignment = Alignment(
                horizontal="center" if i == 1 else "left",
                vertical="center", wrap_text=True)
        ws.row_dimensions[1].height = 24

        HDR        = 5
        hdr_border = Border(top=thick, bottom=thick, left=thin, right=thin)
        for ci, (h, w) in enumerate(zip(self.HEADERS, self.COL_WIDTHS_XLSX), 1):
            c = ws.cell(row=HDR, column=ci, value=h)
            c.font      = Font(name="Arial", bold=True, size=10, color="FFFFFF")
            c.fill      = PatternFill("solid", fgColor=_COLOR_HEADER_BG)
            c.alignment = Alignment(horizontal="center", vertical="center",
                                    wrap_text=True)
            c.border    = hdr_border
            ws.column_dimensions[get_column_letter(ci)].width = w
        ws.row_dimensions[HDR].height = 22

        cell_border = Border(top=thin, bottom=thin, left=thin, right=thin)
        for row in self.rows:
            r    = HDR + row.stt
            c    = _OP_COLORS.get(row.operation, _OP_COLORS["Sửa đổi"])
            fill = PatternFill("solid", fgColor=c["row_fill"])
            data = [row.stt, row.operation, row.scope,
                    row.old_content, row.new_content, row.citation]
            for ci, val in enumerate(data, 1):
                cell = ws.cell(row=r, column=ci, value=str(val) if val else "")
                is_deleted_col = (ci == 4 and row.operation == "Bãi bỏ")
                cell.font = Font(
                    name="Arial", size=9,
                    italic=(ci == 6),
                    color="C00000" if is_deleted_col else "000000",
                )
                cell.fill      = fill
                cell.alignment = Alignment(
                    horizontal="center" if ci in (1, 2) else "left",
                    vertical="top", wrap_text=True)
                cell.border    = cell_border

            if row.highlight_old:
                ws.cell(row=r, column=4).comment = self._comment(
                    f'Cụm từ bị thay: "{row.highlight_old}"')
            if row.highlight_new:
                ws.cell(row=r, column=5).comment = self._comment(
                    f'Cụm từ thay thế: "{row.highlight_new}"')

            ws.row_dimensions[r].height = max(
                40,
                15 * max(1,
                    len(str(row.old_content)) // 40,
                    len(str(row.new_content)) // 40)
            )

        ws.freeze_panes = f"A{HDR + 1}"
        ws.auto_filter.ref = f"A{HDR}:F{HDR + len(self.rows)}"

        self._write_summary_sheet(wb)
        wb.save(output_path)

    def _comment(self, text: str):
        from openpyxl.comments import Comment
        return Comment(text, "Legal Merger")

    def _write_summary_sheet(self, wb):
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        ws   = wb.create_sheet("Tóm tắt")
        thin = Side(style="thin", color="AAAAAA")
        bdr  = Border(top=thin, bottom=thin, left=thin, right=thin)

        ws.merge_cells("A1:C1")
        ws["A1"].value     = "THỐNG KÊ THAY ĐỔI"
        ws["A1"].font      = Font(name="Arial", bold=True, size=12, color="1F4E79")
        ws["A1"].alignment = Alignment(horizontal="center")

        for ci, h in enumerate(["Loại thao tác", "Số lượng", "Màu"], 1):
            c = ws.cell(row=3, column=ci, value=h)
            c.font      = Font(name="Arial", bold=True, size=10, color="FFFFFF")
            c.fill      = PatternFill("solid", fgColor=_COLOR_HEADER_BG)
            c.alignment = Alignment(horizontal="center")
            c.border    = bdr

        counts = {}
        for row in self.rows:
            counts[row.operation] = counts.get(row.operation, 0) + 1

        r = 4
        for op, cnt in counts.items():
            col = _OP_COLORS.get(op, _OP_COLORS["Sửa đổi"])
            ws.cell(row=r, column=1, value=op).font   = Font(name="Arial", size=10)
            ws.cell(row=r, column=1).border = bdr
            ws.cell(row=r, column=2, value=cnt).font  = Font(name="Arial", size=10, bold=True)
            ws.cell(row=r, column=2).alignment = Alignment(horizontal="center")
            ws.cell(row=r, column=2).border = bdr
            ws.cell(row=r, column=3, value=" ").fill   = PatternFill("solid", fgColor=col["label_fill"])
            ws.cell(row=r, column=3).border = bdr
            r += 1

        ws.cell(row=r, column=1, value="Tổng cộng").font = Font(name="Arial", bold=True, size=10)
        ws.cell(row=r, column=1).border = bdr
        ws.cell(row=r, column=2, value=f"=SUM(B4:B{r-1})").font = Font(name="Arial", bold=True)
        ws.cell(row=r, column=2).alignment = Alignment(horizontal="center")
        ws.cell(row=r, column=2).border = bdr

        ws.column_dimensions["A"].width = 20
        ws.column_dimensions["B"].width = 12
        ws.column_dimensions["C"].width = 8
