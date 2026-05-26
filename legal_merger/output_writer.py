"""
output_writer.py — Xuất văn bản hợp nhất ra TXT / DOCX / JSON.
"""

import json
import re
from datetime import datetime

from .models import Node, NodeType

_DOC_TYPE_PREFIX = re.compile(
    r'^(?:Thông tư liên tịch|Thông tư|Nghị định|Quyết định|'
    r'Luật|Pháp lệnh|Chỉ thị|Bộ luật|Hiến pháp|Văn bản)\s+',
    re.IGNORECASE
)

def _short_title(title: str) -> str:
    """Bỏ tiền tố loại văn bản: 'Thông tư 15/2024/TT-NHNN' → '15/2024/TT-NHNN'."""
    return _DOC_TYPE_PREFIX.sub('', title).strip()

# Dòng kết thúc câu hoàn chỉnh (không bị ngắt giữa chừng bởi PDF)
_SENTENCE_END = re.compile(r'[.;!?:»"]\s*$')

# Dòng bắt đầu một phần tử cấu trúc mới (không phải continuation)
_STRUCTURAL_LINE = re.compile(
    r'^\s*(?:Điều\s|\d+[.\)]\s|[a-zđ]\)\s|-\s|'
    r'(?:xx|xix|xviii|xvii|xvi|xv|xiv|xiii|xii|xi|x|ix|viii|vii|vi|v|iv|iii|ii|i)\))',
    re.IGNORECASE
)


_INDENT = {
    NodeType.ARTICLE    : "",
    NodeType.CLAUSE     : "  ",
    NodeType.SUB_CLAUSE : "    ",
    NodeType.POINT      : "      ",
    NodeType.ITEM       : "        ",
}


class OutputWriter:

    # ── TXT ───────────────────────────────────────────────

    def write_txt(self, articles: list, path: str, meta: dict):
        lines = [
            "=" * 65, "VĂN BẢN HỢP NHẤT", "=" * 65,
            f"Văn bản gốc      : {_short_title(meta.get('base_doc',''))}",
            f"Văn bản sửa đổi  : {', '.join(_short_title(d) for d in meta.get('amendment_docs',[]))}",
            f"Ngày hợp nhất: {datetime.now().strftime('%d/%m/%Y')}",
            "=" * 65, "",
        ]
        for art in articles:
            lines += self._render_txt(art)
            lines.append("")
        with open(path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))

    def _render_txt(self, node: Node, depth: int = 0) -> list:
        ind   = _INDENT.get(node.node_type, "  " * depth)
        lines = []

        if node.is_deleted:
            cite = f"  {node.citations[0]}" if node.citations else ""
            lines.append(f"{ind}[BÃI BỎ] {node.label}{cite}")
            return lines

        body = self._merge_wrapped_lines(list(node.body_lines))

        # Label→body continuation (same logic as DOCX renderer)
        label = node.label
        if (body
                and not _SENTENCE_END.search(label)
                and body[0]
                and not _STRUCTURAL_LINE.match(body[0])):
            label = label.rstrip() + ' ' + body[0].lstrip()
            body = body[1:]

        cite_inline = f"  {node.citations[0]}" if node.citations else ""
        lines.append(f"{ind}{label}{cite_inline}")

        for c in node.citations[1:]:
            lines.append(f"{ind}  {c}")

        for bl in body:
            lines.append(f"{ind}{bl}")

        for child in node.children:
            lines += self._render_txt(child, depth + 1)

        return lines

    # ── DOCX ──────────────────────────────────────────────

    def write_docx(self, articles: list, path: str, meta: dict):
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        for section in doc.sections:
            section.top_margin = section.bottom_margin = Cm(2)
            section.left_margin = section.right_margin = Cm(3)

        from docx.enum.text import WD_LINE_SPACING

        normal = doc.styles['Normal']
        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(14)
        normal.paragraph_format.space_before       = Pt(6)
        normal.paragraph_format.space_after        = Pt(6)
        normal.paragraph_format.alignment          = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.first_line_indent  = Cm(1.25)
        normal.paragraph_format.line_spacing_rule  = WD_LINE_SPACING.MULTIPLE
        normal.paragraph_format.line_spacing       = 1.15

        t = doc.add_heading("VĂN BẢN HỢP NHẤT", level=1)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        r = p.add_run("Văn bản gốc: ")
        r.bold = True
        p.add_run(_short_title(meta.get('base_doc', '')))
        p2 = doc.add_paragraph()
        r2 = p2.add_run("Văn bản sửa đổi: ")
        r2.bold = True
        p2.add_run(', '.join(_short_title(d) for d in meta.get('amendment_docs', [])))
        doc.add_paragraph(f"Ngày hợp nhất: {datetime.now().strftime('%d/%m/%Y')}")
        doc.add_paragraph()

        for art in articles:
            self._render_docx(doc, art)

        doc.save(path)

    def _render_docx(self, doc, node: Node, depth: int = 0):
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

        CITE_COLOR = RGBColor(0x44, 0x72, 0xC4)
        DEL_COLOR  = RGBColor(0x99, 0x99, 0x99)

        def _fmt(p):
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.left_indent       = Cm(0)
            p.paragraph_format.space_before      = Pt(6)
            p.paragraph_format.space_after       = Pt(6)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing      = 1.15
            p.alignment                          = WD_ALIGN_PARAGRAPH.JUSTIFY

        if node.is_deleted:
            p   = doc.add_paragraph()
            _fmt(p)
            run = p.add_run(f"[BÃI BỎ] {node.label}")
            run.font.color.rgb = DEL_COLOR
            run.font.strike    = True
            run.font.size      = Pt(10)
            if node.citations:
                cr = p.add_run(f"  {node.citations[0]}")
                cr.font.color.rgb = DEL_COLOR
                cr.font.size      = Pt(9)
                cr.font.italic    = True
            return

        p = doc.add_paragraph()
        _fmt(p)

        run = p.add_run(node.label)
        run.bold       = (node.node_type == NodeType.ARTICLE)
        run.font.name  = 'Times New Roman'
        run.font.size  = Pt(14)

        body = self._merge_wrapped_lines(list(node.body_lines))

        # Label→body continuation: nếu label kết thúc giữa câu và body_line đầu
        # tiên là tiếp nối (chữ thường, không phải phần tử cấu trúc) → nối vào
        # cùng một đoạn với label thay vì tạo paragraph mới.
        if (body
                and not _SENTENCE_END.search(node.label)
                and body[0]
                and not _STRUCTURAL_LINE.match(body[0])):
            cont = p.add_run(' ' + body[0].lstrip())
            cont.font.name = 'Times New Roman'
            cont.font.size = Pt(14)
            body = body[1:]

        last_para = p
        for bl in body:
            bp = doc.add_paragraph(bl)
            _fmt(bp)
            for r in bp.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(14)
            last_para = bp

        if node.citations:
            for cite in node.citations:
                cr = last_para.add_run(f"  {cite}")
                cr.font.color.rgb = CITE_COLOR
                cr.font.size      = Pt(11)
                cr.font.italic    = True
                cr.font.name      = 'Times New Roman'

        for child in node.children:
            self._render_docx(doc, child, depth + 1)

    # ── Helpers ───────────────────────────────────────────

    @staticmethod
    def _merge_wrapped_lines(lines: list) -> list:
        """Nối các dòng bị PDF ngắt giữa câu thành một dòng logic.

        Một dòng được coi là tiếp nối (continuation) nếu:
        - Dòng trước KHÔNG kết thúc bằng dấu câu kết thúc (.;!?:»")
        - Dòng kế bắt đầu bằng chữ thường (không phải phần tử cấu trúc)
        """
        if not lines:
            return lines
        merged = [lines[0]]
        for line in lines[1:]:
            prev = merged[-1]
            if (prev
                    and not _SENTENCE_END.search(prev)
                    and line
                    and not _STRUCTURAL_LINE.match(line)):
                merged[-1] = prev.rstrip() + ' ' + line.lstrip()
            else:
                merged.append(line)
        return merged

    # ── Báo cáo JSON ──────────────────────────────────────

    def write_change_report(self, merge_log: list, path: str):
        report = {
            "generated_at" : datetime.now().isoformat(),
            "total_changes": len([e for e in merge_log if e.get("action") != "CẢNH BÁO"]),
            "warnings"     : len([e for e in merge_log if e.get("action") == "CẢNH BÁO"]),
            "changes"      : merge_log,
        }
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)
